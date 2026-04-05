"""
viewers/text_viewer.py - Viewer for TXT, Markdown, and DOCX files.
Uses QTextBrowser for rich display with text selection support.
"""
import os
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QTextBrowser,
    QPushButton, QLabel, QSizePolicy, QSpacerItem,
    QFrame, QSlider
)
from PyQt5.QtCore import Qt, pyqtSignal
from PyQt5.QtGui import QFont, QTextCursor


class TextViewer(QWidget):
    """Viewer for TXT, Markdown, and DOCX files with text selection support."""
    selection_changed = pyqtSignal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._current_format = "txt"
        self._font_size = 13
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # ── Toolbar ──────────────────────────────────────────────────────────
        toolbar = QFrame()
        toolbar.setFixedHeight(40)
        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(8, 4, 8, 4)
        tb_layout.setSpacing(6)

        self.lbl_info = QLabel("Sin documento")
        self.lbl_info.setObjectName("subtitle")

        btn_smaller = QPushButton("A−")
        btn_smaller.setFixedWidth(36)
        btn_smaller.setToolTip("Reducir texto")
        btn_smaller.clicked.connect(self.decrease_font)

        btn_larger = QPushButton("A+")
        btn_larger.setFixedWidth(36)
        btn_larger.setToolTip("Aumentar texto")
        btn_larger.clicked.connect(self.increase_font)

        btn_copy_sel = QPushButton("📋 Copiar selección")
        btn_copy_sel.setFixedWidth(130)
        btn_copy_sel.setToolTip("Copiar texto seleccionado")
        btn_copy_sel.clicked.connect(self.copy_selection)

        tb_layout.addWidget(self.lbl_info)
        tb_layout.addItem(QSpacerItem(0, 0, QSizePolicy.Expanding, QSizePolicy.Minimum))
        tb_layout.addWidget(btn_smaller)
        tb_layout.addWidget(btn_larger)
        tb_layout.addWidget(btn_copy_sel)

        layout.addWidget(toolbar)

        # ── Text browser ──────────────────────────────────────────────────────
        self.browser = QTextBrowser()
        self.browser.setOpenExternalLinks(False)
        self.browser.setReadOnly(True)
        self.browser.selectionChanged.connect(self._on_selection_changed)
        layout.addWidget(self.browser)

        self._apply_font()

    def _apply_font(self):
        font = QFont("Segoe UI", self._font_size)
        self.browser.setFont(font)

    def load_document(self, path):
        """Load a document from path based on its extension."""
        ext = os.path.splitext(path)[1].lower()
        self._current_format = ext.lstrip(".")
        name = os.path.basename(path)
        self.lbl_info.setText(f"📄 {name}")

        if ext in (".txt",):
            self._load_txt(path)
        elif ext in (".md",):
            self._load_md(path)
        elif ext in (".docx",):
            self._load_docx(path)
        else:
            self._load_txt(path)

    def _load_txt(self, path):
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            self.browser.setFont(QFont("Courier New", self._font_size))
            # Wrap in basic HTML for nicer display
            escaped = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html = f"""<html><body style="font-family: 'Courier New'; font-size: {self._font_size}pt;
                        white-space: pre-wrap; line-height: 1.6; padding: 20px;">
                        {escaped}</body></html>"""
            self.browser.setHtml(html)
        except Exception as e:
            self.browser.setPlainText(f"Error al leer el archivo:\n{str(e)}")

    def _load_md(self, path):
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            try:
                import markdown
                html_body = markdown.markdown(
                    content,
                    extensions=["fenced_code", "tables", "toc", "nl2br"]
                )
            except ImportError:
                # Fallback: basic formatting
                escaped = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                html_body = f"<pre>{escaped}</pre>"

            html = f"""<html><head><style>
                body {{ font-family: 'Segoe UI', sans-serif; font-size: {self._font_size}pt;
                        line-height: 1.7; padding: 24px; max-width: 860px; margin: 0 auto; }}
                h1,h2,h3,h4 {{ color: #89b4fa; margin-top: 1.2em; }}
                h1 {{ border-bottom: 2px solid #89b4fa; padding-bottom: 6px; }}
                h2 {{ border-bottom: 1px solid #45475a; padding-bottom: 4px; }}
                code {{ background: #313244; padding: 2px 6px; border-radius: 4px;
                        font-family: 'Courier New'; font-size: 90%; }}
                pre {{ background: #181825; padding: 14px; border-radius: 8px;
                       border-left: 3px solid #89b4fa; overflow-x: auto; }}
                pre code {{ background: none; padding: 0; }}
                blockquote {{ border-left: 4px solid #89b4fa; margin-left: 0;
                              padding-left: 16px; color: #6c7086; }}
                table {{ border-collapse: collapse; width: 100%; }}
                th {{ background: #313244; padding: 8px 12px; text-align: left;
                      border: 1px solid #45475a; }}
                td {{ padding: 8px 12px; border: 1px solid #313244; }}
                tr:nth-child(even) {{ background: #181825; }}
                a {{ color: #89b4fa; }}
                hr {{ border: none; border-top: 1px solid #313244; }}
                </style></head><body>{html_body}</body></html>"""
            self.browser.setHtml(html)
        except Exception as e:
            self.browser.setPlainText(f"Error al leer el archivo:\n{str(e)}")

    def _load_docx(self, path):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            doc = Document(path)

            html_parts = [f"""<html><head><style>
                body {{ font-family: 'Segoe UI', sans-serif; font-size: {self._font_size}pt;
                        line-height: 1.7; padding: 32px; max-width: 860px; margin: 0 auto; }}
                h1 {{ color: #89b4fa; font-size: 1.8em; margin-top: 0.5em; }}
                h2 {{ color: #89b4fa; font-size: 1.4em; }}
                h3 {{ color: #cdd6f4; font-size: 1.2em; }}
                table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
                td, th {{ border: 1px solid #45475a; padding: 8px 12px; }}
                th {{ background: #313244; }}
                .bold {{ font-weight: bold; }}
                .italic {{ font-style: italic; }}
                .underline {{ text-decoration: underline; }}
                </style></head><body>"""]

            for para in doc.paragraphs:
                style_name = para.style.name.lower()
                text = ""
                for run in para.runs:
                    rt = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    if run.bold:
                        rt = f"<b>{rt}</b>"
                    if run.italic:
                        rt = f"<i>{rt}</i>"
                    if run.underline:
                        rt = f"<u>{rt}</u>"
                    text += rt

                if not text.strip():
                    html_parts.append("<p>&nbsp;</p>")
                elif "heading 1" in style_name:
                    html_parts.append(f"<h1>{text}</h1>")
                elif "heading 2" in style_name:
                    html_parts.append(f"<h2>{text}</h2>")
                elif "heading 3" in style_name:
                    html_parts.append(f"<h3>{text}</h3>")
                elif "list" in style_name:
                    html_parts.append(f"<li>{text}</li>")
                else:
                    html_parts.append(f"<p>{text}</p>")

            # Tables
            for table in doc.tables:
                html_parts.append("<table>")
                for i, row in enumerate(table.rows):
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        tag = "th" if i == 0 else "td"
                        cell_text = cell.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        html_parts.append(f"<{tag}>{cell_text}</{tag}>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")

            html_parts.append("</body></html>")
            self.browser.setHtml("".join(html_parts))

        except ImportError:
            self.browser.setPlainText(
                "python-docx no está instalado.\nEjecuta: pip install python-docx"
            )
        except Exception as e:
            self.browser.setPlainText(f"Error al leer el archivo DOCX:\n{str(e)}")

    def get_selected_text(self):
        cursor = self.browser.textCursor()
        return cursor.selectedText()

    def get_text_content(self):
        """Return plain text content for full-text search indexing."""
        return self.browser.toPlainText()

    def _on_selection_changed(self):
        text = self.get_selected_text()
        if text:
            self.selection_changed.emit(text)

    def copy_selection(self):
        from PyQt5.QtWidgets import QApplication
        text = self.get_selected_text()
        if text:
            QApplication.clipboard().setText(text)

    def increase_font(self):
        self._font_size = min(24, self._font_size + 1)
        self._apply_font()

    def decrease_font(self):
        self._font_size = max(8, self._font_size - 1)
        self._apply_font()
