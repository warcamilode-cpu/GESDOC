"""
viewers/web_viewer.py

SmartViewer routing:
  PDF           → PDFViewer (PyMuPDF image renderer) — stable, threaded
  MD/DOCX/TXT   → WebEngineViewer (Chromium HTML)    — stable for HTML
"""
import os
import tempfile
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QLabel, QLineEdit, QFrame, QSizePolicy
)
from PyQt5.QtCore import Qt, pyqtSignal, QUrl, QTimer
from PyQt5.QtGui import QFont

# ── Try WebEngine (only used for non-PDF formats) ─────────────────────────────
try:
    from PyQt5.QtWebEngineWidgets import QWebEngineView, QWebEngineSettings, QWebEnginePage

    class _SilentPage(QWebEnginePage):
        _IGNORE = ("Permission denied", "Unchecked runtime", "Cannot read property",
                   "Assertion failed", "getStrings", "metricsPrivate")
        def javaScriptConsoleMessage(self, level, msg, line, src):
            if any(k in msg for k in self._IGNORE):
                return
            super().javaScriptConsoleMessage(level, msg, line, src)

    HAS_WEBENGINE = True
except Exception as _e:
    print(f"[DocManager] PyQtWebEngine no disponible: {_e}")
    HAS_WEBENGINE = False


# ─────────────────────────────────────────────────────────────────────────────
# HTML converters for non-PDF formats
# ─────────────────────────────────────────────────────────────────────────────

def _md_to_html(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    try:
        import markdown
        body = markdown.markdown(content, extensions=["fenced_code", "tables", "toc", "nl2br"])
    except ImportError:
        esc = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        body = f"<pre style='white-space:pre-wrap'>{esc}</pre>"
    return _wrap_html(body, os.path.basename(path))


def _txt_to_html(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    esc = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    body = f"<pre style='white-space:pre-wrap;font-family:\"Courier New\",monospace'>{esc}</pre>"
    return _wrap_html(body, os.path.basename(path))


def _docx_to_html(path):
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            style = para.style.name.lower()
            text = ""
            for run in para.runs:
                rt = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if run.bold:      rt = f"<strong>{rt}</strong>"
                if run.italic:    rt = f"<em>{rt}</em>"
                if run.underline: rt = f"<u>{rt}</u>"
                text += rt
            if not text.strip():
                parts.append("<p>&nbsp;</p>")
            elif "heading 1" in style: parts.append(f"<h1>{text}</h1>")
            elif "heading 2" in style: parts.append(f"<h2>{text}</h2>")
            elif "heading 3" in style: parts.append(f"<h3>{text}</h3>")
            elif "list"      in style: parts.append(f"<li>{text}</li>")
            else:                       parts.append(f"<p>{text}</p>")
        for table in doc.tables:
            parts.append("<table>")
            for i, row in enumerate(table.rows):
                parts.append("<tr>")
                for cell in row.cells:
                    tag = "th" if i == 0 else "td"
                    ct = cell.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    parts.append(f"<{tag}>{ct}</{tag}>")
                parts.append("</tr>")
            parts.append("</table>")
        return _wrap_html("".join(parts), os.path.basename(path))
    except ImportError:
        return _wrap_html("<p style='color:#f38ba8'>Instala python-docx: <code>pip install python-docx</code></p>", "Error")
    except Exception as e:
        return _wrap_html(f"<p style='color:#f38ba8'>Error: {e}</p>", "Error")


def _wrap_html(body, title=""):
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8"><title>{title}</title>
<style>
  :root{{--bg:#1e1e2e;--surface:#181825;--text:#cdd6f4;--muted:#6c7086;
         --accent:#89b4fa;--code-bg:#313244;--border:#45475a;}}
  *{{box-sizing:border-box;margin:0;padding:0;}}
  body{{font-family:'Times New Roman',Times,serif;font-size:14px;line-height:1.75;
        color:var(--text);background:var(--bg);max-width:860px;margin:0 auto;
        padding:36px 40px 60px;}}
  h1,h2,h3{{color:var(--accent);margin:1.3em 0 .5em;font-weight:600;}}
  h1{{font-size:1.85em;border-bottom:2px solid var(--accent);padding-bottom:8px;}}
  h2{{font-size:1.4em;border-bottom:1px solid var(--border);padding-bottom:5px;}}
  p{{margin:.7em 0;}} a{{color:var(--accent);}}
  strong{{color:#cba6f7;}} em{{color:#fab387;font-style:italic;}}
  pre{{font-family:'Cascadia Code','Courier New',monospace;font-size:13px;
       background:var(--surface);border-left:4px solid var(--accent);
       border-radius:8px;padding:16px 20px;overflow-x:auto;margin:1em 0;}}
  code{{font-family:'Cascadia Code','Courier New',monospace;font-size:89%;
        background:var(--code-bg);color:#a6e3a1;padding:2px 6px;border-radius:4px;}}
  pre code{{background:none;padding:0;color:var(--text);}}
  blockquote{{border-left:4px solid var(--accent);background:var(--surface);
              margin:1em 0;padding:10px 18px;border-radius:0 6px 6px 0;color:var(--muted);}}
  table{{border-collapse:collapse;width:100%;margin:1.2em 0;}}
  th{{background:var(--code-bg);color:var(--accent);padding:10px 14px;
      text-align:left;border:1px solid var(--border);font-weight:600;}}
  td{{padding:8px 14px;border:1px solid var(--border);}}
  tr:nth-child(even){{background:var(--surface);}}
  li{{margin:.3em 0 .3em 1.6em;}} ul,ol{{margin:.6em 0;}}
</style></head><body>
{body}
<script>
document.addEventListener('mouseup',function(){{
  var s=window.getSelection().toString().trim();
  if(s)window._dmSelection=s;
}});
</script>
</body></html>"""


# ─────────────────────────────────────────────────────────────────────────────
# WebEngineViewer — only for HTML/MD/DOCX/TXT
# ─────────────────────────────────────────────────────────────────────────────

class WebEngineViewer(QWidget):
    """Chromium-based viewer for text documents (MD, DOCX, TXT)."""
    selection_changed = pyqtSignal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._tmp_files = []
        self._current_path = ""
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Toolbar
        bar = QFrame()
        bar.setFixedHeight(38)
        bar.setStyleSheet("QFrame{background:#181825;border-bottom:1px solid #313244;}")
        tb = QHBoxLayout(bar)
        tb.setContentsMargins(10, 4, 10, 4)
        tb.setSpacing(8)

        self.lbl_info = QLabel("Sin documento")
        self.lbl_info.setStyleSheet(
            "color:#6c7086;font-family:'Times New Roman';font-size:14px;background:transparent;")

        lbl_engine = QLabel("🌐 Chrome engine")
        lbl_engine.setStyleSheet(
            "color:#a6e3a1;font-family:'Times New Roman';font-size:14px;background:transparent;")

        self.btn_capture = QPushButton("🖊 Capturar selección")
        self.btn_capture.setFixedHeight(26)
        self.btn_capture.setToolTip("Selecciona texto y presiona para vincularlo a un comentario")
        self.btn_capture.clicked.connect(self._capture_selection)

        tb.addWidget(self.lbl_info)
        tb.addStretch()
        tb.addWidget(lbl_engine)
        tb.addWidget(self.btn_capture)
        layout.addWidget(bar)

        # Web view
        self.web = QWebEngineView()
        self.web.setPage(_SilentPage(self.web))
        s = self.web.settings()
        s.setAttribute(QWebEngineSettings.LocalContentCanAccessFileUrls, True)
        s.setAttribute(QWebEngineSettings.JavascriptEnabled,             True)
        s.setAttribute(QWebEngineSettings.ScrollAnimatorEnabled,         True)
        layout.addWidget(self.web)

    def load_document(self, path):
        self._current_path = path
        name = os.path.basename(path)
        self.lbl_info.setText(f"📄 {name}")
        fmt = os.path.splitext(path)[1].lstrip(".").lower()

        if fmt == "md":
            self._load_html(_md_to_html(path))
        elif fmt in ("docx", "doc"):
            self._load_html(_docx_to_html(path))
        else:
            self._load_html(_txt_to_html(path))

    def get_selected_text(self):
        return getattr(self, "_last_selection", "")

    def get_text_content(self):
        if not self._current_path:
            return ""
        fmt = os.path.splitext(self._current_path)[1].lstrip(".").lower()
        try:
            if fmt in ("md", "txt"):
                return open(self._current_path, "r", encoding="utf-8", errors="replace").read()
            elif fmt in ("docx", "doc"):
                from docx import Document
                return "\n".join(p.text for p in Document(self._current_path).paragraphs)
        except Exception:
            pass
        return ""

    def get_current_page_info(self):
        return f"📄 {os.path.basename(self._current_path)}" if self._current_path else ""

    def _load_html(self, html):
        tmp = tempfile.NamedTemporaryFile(
            mode="w", suffix=".html", encoding="utf-8",
            delete=False, prefix="docmanager_"
        )
        tmp.write(html)
        tmp.close()
        self._tmp_files.append(tmp.name)
        self.web.load(QUrl.fromLocalFile(tmp.name))

    def _capture_selection(self):
        self.web.page().runJavaScript(
            "window.getSelection().toString();",
            self._on_selection_captured
        )

    def _on_selection_captured(self, text):
        if text:
            self._last_selection = text.strip()
            self.selection_changed.emit(self._last_selection)
            preview = text[:60] + ("..." if len(text) > 60 else "")
            self.lbl_info.setText(f"✅ «{preview}»")

    def closeEvent(self, event):
        for f in self._tmp_files:
            try:
                os.remove(f)
            except Exception:
                pass
        super().closeEvent(event)


# ─────────────────────────────────────────────────────────────────────────────
# SmartViewer — routes by format
# ─────────────────────────────────────────────────────────────────────────────

class SmartViewer(QWidget):
    """
    Routes documents to the best viewer:
      PDF  → PDFViewer (PyMuPDF, always stable)
      else → WebEngineViewer (Chromium HTML, if available) or TextViewer fallback
    """
    selection_changed = pyqtSignal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._viewer = None
        self._layout = QVBoxLayout(self)
        self._layout.setContentsMargins(0, 0, 0, 0)
        self._layout.setSpacing(0)

    def load_document(self, path):
        fmt = os.path.splitext(path)[1].lstrip(".").lower()

        # Tear down previous viewer
        if self._viewer is not None:
            self._layout.removeWidget(self._viewer)
            self._viewer.hide()
            self._viewer.deleteLater()
            self._viewer = None

        if fmt == "pdf":
            # Always use PyMuPDF — reliable, threaded, no Chromium PDF issues
            from viewers.pdf_viewer import PDFViewer
            v = PDFViewer()
            v.load_document(path)

        elif HAS_WEBENGINE:
            # Use Chromium for text-based formats
            v = WebEngineViewer()
            v.load_document(path)

        else:
            # Full fallback
            from viewers.text_viewer import TextViewer
            v = TextViewer()
            v.load_document(path)

        if hasattr(v, "selection_changed"):
            v.selection_changed.connect(self.selection_changed)

        self._viewer = v
        self._layout.addWidget(v)
        v.show()

    def get_selected_text(self):
        if self._viewer and hasattr(self._viewer, "get_selected_text"):
            return self._viewer.get_selected_text()
        return ""

    def get_text_content(self):
        if self._viewer and hasattr(self._viewer, "get_text_content"):
            return self._viewer.get_text_content()
        return ""

    def get_current_page_info(self):
        if self._viewer and hasattr(self._viewer, "get_current_page_info"):
            return self._viewer.get_current_page_info()
        return ""

    def request_text_content_async(self, callback):
        """For compatibility with main_window._index_content."""
        content = self.get_text_content()
        if content:
            callback(content)
